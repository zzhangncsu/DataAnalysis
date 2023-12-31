import requests
import streamlit as st
import pandas as pd
import numpy as np

import streamlit as st
import pandas as pd
from docx import Document
import io
from docx.shared import RGBColor
import datetime
import urllib
from collections import defaultdict
import streamlit.components.v1 as components

import pathlib
from bs4 import BeautifulSoup
import shutil
import logging

if 'key' not in st.session_state:
    st.session_state['key'] = 'value'

day_mappings = {'Sunday': "周日",
                'Monday': "周一",
                'Tuesday': "周二",
                'Wednesday': "周三",
                'Thursday': "周四",
                'Friday': "周五",
                'Saturday': "周六", }


def inject_ga():
    GA_ID = "google_analytics"

    GA_JS = """
    <!-- Google tag (gtag.js) -->
<script async src="https://www.googletagmanager.com/gtag/js?id=G-005TYTNJ4W"></script>
<script>
  window.dataLayer = window.dataLayer || [];
  function gtag(){dataLayer.push(arguments);}
  gtag('js', new Date());

  gtag('config', 'G-005TYTNJ4W');
</script>
    """

    # Insert the script in the head tag of the static template inside your virtual
    index_path = pathlib.Path(st.__file__).parent / "static" / "index.html"
    logging.info(f's {index_path}')
    soup = BeautifulSoup(index_path.read_text(), features="html.parser")
    if not soup.find(id=GA_ID):
        # bck_index = index_path.with_suffix('.bck')
        # if bck_index.exists():
        #     shutil.copy(bck_index, index_path)
        # else:
        #     shutil.copy(index_path, bck_index)
        html = str(soup)
        new_html = html.replace('<head>', '<head>\n' + GA_JS)
        index_path.write_text(new_html)


def getAdsSeries(text):
    ads = text[:text.rfind('-')]
    last = text[text.rfind('-') + 1:]
    if last in ['视频', '轮播', '绿白黑视频']:
        text = ads
        ads = text[:text.rfind('-')]
        last = text[text.rfind('-') + 1:]
    return ads, last


#
# inject_ga()

uploaded_file = st.file_uploader("Choose a file")
if uploaded_file is not None:
    df = pd.read_csv(uploaded_file)

    all_date = sorted(df['day'].unique())
    df = df[df['utm_campaign_content'].notna()]
    df['utm_campaign_content'] = df['utm_campaign_content'].apply(lambda x: urllib.parse.unquote(x))
    df = df[['utm_campaign_content', 'day', 'total_carts', 'total_checkouts', 'total_orders_placed']]
    df['ads'] = df['utm_campaign_content'].apply(lambda x: getAdsSeries(x)[0])
    df['series'] = df['utm_campaign_content'].apply(lambda x: getAdsSeries(x)[1])

    # remove untracked data
    df = df[df['ads'] != 'Facebook_U']
    df = df[df['ads'] != 'Shoppingpmax']
    df = df[df['ads'] != 'sag_organi']

    # get data with at least one type of activities
    df_focus = df.loc[~((df['total_carts'] == 0) & (df['total_checkouts'] == 0) & (df['total_orders_placed'] == 0))]

    ads_groups = df.groupby(['ads'])

    ads_dict = defaultdict(list)
    for ads in ads_groups.groups:
        ads_dict[ads] = ads_groups.get_group(ads)['series'].unique()

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["🐰", "文档", "统计", "图标", "SHEIN"])
    with tab2:
        document = Document()
        document.add_heading('数据', 0)

        for ads in ads_groups.groups:
            if ads in df_focus['ads'].unique():
                pm = document.add_paragraph()
                rm = pm.add_run(f'{ads}, {len(ads_dict[ads])}个标签')
                rm.add_break()
                tags = ', '.join(ads_dict[ads])
                rm = pm.add_run(f'标签: {tags}')
                rm.add_break()
                ads_focus = df_focus[df_focus['ads'] == ads]
                for date in all_date:
                    if date in ads_focus['day'].unique():
                        date_focus = ads_focus[ads_focus['day'] == date]
                        value = []
                        date_focus = date_focus.sort_values(by=['total_orders_placed'], ascending=False)
                        day = datetime.datetime.strptime(date, '%Y-%m-%d').strftime('%A')
                        rm = pm.add_run(f"{date} ({day_mappings[day]}): ")
                        count = 0
                        run = None
                        for idx, row in date_focus.iterrows():
                            run = pm.add_run(
                                f"{row['series']} {row['total_carts']}-{row['total_checkouts']}-{row['total_orders_placed']}")
                            if row['total_orders_placed'] > 0:
                                font = run.font
                                font.color.rgb = RGBColor(10, 150, 10)
                            if count != len(date_focus) - 1:
                                pm.add_run(", ")
                            count += 1
                        if run is not None:
                            run.add_break()

        p = document.add_paragraph()
        run = p.add_run("喜欢你 兔兔")
        font = run.font
        font.color.rgb = RGBColor(220, 220, 220)

        file_stream = io.BytesIO()
        document.save(file_stream)
        st.download_button('下载', file_stream, file_name='data.docx')

    with tab3:
        aggre_dict = defaultdict(defaultdict)
        for ads in ads_groups.groups:
            aggre_dict[ads]['total_carts'] = ads_groups.get_group(ads)['total_carts'].sum()
            aggre_dict[ads]['total_checkouts'] = ads_groups.get_group(ads)['total_checkouts'].sum()
            aggre_dict[ads]['total_orders_placed'] = ads_groups.get_group(ads)['total_orders_placed'].sum()
            aggre_dict[ads]['days'] = len(ads_groups.get_group(ads)['day'].unique())
            aggre_dict[ads]['tags'] = ads_groups.get_group(ads)['series'].unique()
        aggre_df = pd.DataFrame.from_dict(aggre_dict).T
        st.dataframe(aggre_df)

        option = st.selectbox(
            '广告系列', aggre_df.index)

        aggre_date_dict = defaultdict(lambda: defaultdict(dict))
        for ads in ads_groups.groups:
            ads_day_groups = ads_groups.get_group(ads).groupby('day')
            for day in ads_day_groups.groups:
                aggre_date_dict[ads][day]['total_carts'] = ads_day_groups.get_group(day)['total_carts'].sum()
                aggre_date_dict[ads][day]['total_checkouts'] = ads_day_groups.get_group(day)['total_checkouts'].sum()
                aggre_date_dict[ads][day]['total_orders_placed'] = ads_day_groups.get_group(day)[
                    'total_orders_placed'].sum()
        df_prod = pd.DataFrame.from_dict(aggre_date_dict[option]).T

        df_group_tag = ads_groups.get_group(option).groupby(['series']).sum(
            ['total_carts', 'total_checkouts', 'total_orders_placed'])
        df_group_tag.index.names = ['Tags']
        st.dataframe(df_group_tag)
        st.dataframe(df_prod)
        #
        st.bar_chart(df_prod)
    with tab4:
        st.image("https://cdn.midjourney.com/1eea2b76-9e24-4a3f-ac81-ebdb7fd84389/0_2.png")
        st.image("https://cdn.midjourney.com/0eed04e8-49da-4b19-9701-460f7824371c/0_0.png")
        st.image("https://blog.lvhglobal.com/content/images/2018/10/hawaii-sunset-main.png")
        st.image("https://cdn.midjourney.com/608bee51-598d-48bb-a6ac-3276c41b0f7c/0_1.png")
        st.image(
            "https://scontent-sjc3-1.cdninstagram.com/v/t39.30808-6/336328467_544655537773767_7949993889363802673_n.jpg?stp=c0.60.1440.1800a_dst-jpg_e15&efg=eyJ2ZW5jb2RlX3RhZyI6ImltYWdlX3VybGdlbi4xNDQweDE5MjAuc2RyIn0&_nc_ht=scontent-sjc3-1.cdninstagram.com&_nc_cat=103&_nc_ohc=wpPa2yr4IA8AX-IILoh&edm=ABmJApAAAAAA&ccb=7-5&ig_cache_key=MzA2MDE3NTU4OTQ5NTA3NDEwMw%3D%3D.2-ccb7-5&oh=00_AfBeVTEPEtS5qTEjhers68r9Tzqhmhbb8hSPf_IPV_SPlw&oe=6525C0AE&_nc_sid=b41fef")

        # components.iframe("plot.html")
    with tab1:
        components.html(
            """
              <style>
              :root {
          --primary-color: #212121;
          --background-color: #111;
          --font: sans-serif;
        }
        
        * {
          margin: 0;
          padding: 0;
        }
        
        body {
          background: var(--background-color);
          font-family: var(--font);
          display: flex;
          justify-content: center;
        }
        
        /* Timeline Container */
        .timeline {
          background: var(--primary-color);
          margin: 20px auto;
          padding: 20px;
        }
        
        /* Card container */
        .card {
          position: relative;
          max-width: 600px;
        }
        
        /* setting padding based on even or odd */
        .card:nth-child(odd) {
          padding: 30px 0 30px 30px;
        }
        .card:nth-child(even) {
          padding: 30px 30px 30px 0;
        }
        /* Global ::before */
        .card::before {
          content: "";
          position: absolute;
          width: 50%;
          border: solid orangered;
        }
        
        /* Setting the border of top, bottom, left */
        .card:nth-child(odd)::before {
          left: 0px;
          top: -4.5px;
          bottom: -4.5px;
          border-width: 5px 0 5px 5px;
          border-radius: 50px 0 0 50px;
        }
        
        /* Setting the border of top, bottom, right */
        .card:nth-child(even)::before {
          right: 0;
          top: 0;
          bottom: 0;
          border-width: 5px 5px 5px 0;
          border-radius: 0 50px 50px 0;
        }
        
        /* Removing the border if it is the first card */
        .card:first-child::before {
          border-top: 0;
          border-top-left-radius: 0;
        }
        
        /* Removing the border if it is the last card  and it's odd */
        .card:last-child:nth-child(odd)::before {
          border-bottom: 0;
          border-bottom-left-radius: 0;
        }
        
        /* Removing the border if it is the last card  and it's even */
        .card:last-child:nth-child(even)::before {
          border-bottom: 0;
          border-bottom-right-radius: 0;
        }
        
        /* Information about the timeline */
        .info {
          display: flex;
          flex-direction: column;
          background: #333;
          color: gray;
          border-radius: 10px;
          padding: 10px;
        }
        
        /* Title of the card */
        .title {
          color: orangered;
          position: relative;
        }
        
        /* Timeline dot  */
        .title::before {
          content: "";
          position: absolute;
          width: 10px;
          height: 10px;
          background: white;
          border-radius: 999px;
          border: 3px solid orangered;
        }
        
        /* text right if the card is even  */
        .card:nth-child(even) > .info > .title {
          text-align: right;
        }
        
        /* setting dot to the left if the card is odd */
        .card:nth-child(odd) > .info > .title::before {
          left: -45px;
        }
        
        /* setting dot to the right if the card is odd */
        .card:nth-child(even) > .info > .title::before {
          right: -45px;
        }
        
        a:link {
          color: orangered;
          background-color: transparent;
          text-decoration: none;
        }
        
        a:visited {
          color: orangered;
          background-color: transparent;
          text-decoration: none;
        }
        
        a:hover {
          color: orangered;
          background-color: transparent;
          text-decoration: underline;
        }
        
              </style>
            <div class="timeline">
        <div class="outer">
                                <div class="card">
        <div class="info">
        <h3 class="title">Oct 19, 2023 星期四 南充</h3>
                <p><img src="https://cdn.midjourney.com/9aac7f61-ab3c-40b3-bc5d-374366a89aea/0_0.png"  height="500" align="left"></p>
        </div>
        </div>
                        <div class="card">
        <div class="info">
        <h3 class="title">涨!</h3>
                <p><img src="https://raw.githubusercontent.com/zzhangncsu/DataAnalysis/main/hand.jpg"  height="300" align="right"></p>
        </div>
        </div>
                <div class="card">
        <div class="info">
        <h3 class="title">兔兔不去 😭 </h3>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">10/28 周六 北京 到 檀香山 </h3>
        <p><br><img src="https://raw.githubusercontent.com/zzhangncsu/DataAnalysis/main/1.png"  height="250"></p>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">10/29 周日 Wakiki Beach 购物 看海 看日落!</h3>
        <p><br><img src="https://raw.githubusercontent.com/zzhangncsu/DataAnalysis/main/2.jpg"  height="300" align="right"></p>
        <p><br><img src="https://raw.githubusercontent.com/zzhangncsu/DataAnalysis/main/3.jpg"  height="250" align="right"></p>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">10/30 周一 可爱岛 君悦 DeHome Workshop Day 1!</h3>
        <p><br><img src="https://raw.githubusercontent.com/zzhangncsu/DataAnalysis/main/4.jpg"  height="250"></p>
        <p><br><img src="https://raw.githubusercontent.com/zzhangncsu/DataAnalysis/main/dehome.png"  height="250"></p>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">10/31 周二 可爱岛 Westin Princeville 万圣节！</h3>
        <p><br><a href="https://www.marriott.com/en-us/hotels/lihwp-the-westin-princeville-ocean-resort-villas/overview/"><img src="https://imgcy.trivago.com/c_limit,d_dummy.jpeg,f_auto,h_600,q_auto,w_600//hotelier-images/50/f6/2aabc8771417b888f44471ce1d469121d47879ee0c64bc16a017291617ab.jpeg"  height="250" align="right"></a></p>
        <p><br><img src="https://rideholoholo.com/wp-content/uploads/2022/10/halloween-post.jpg"  height="250" align="right"></p>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">11/01 周三 可爱岛 一起上班 忙工作 喝咖啡 吃好吃的</h3>
        <p><br><img src="https://hips.hearstapps.com/hmg-prod/images/workfromhawaii-1526404272.jpeg"  height="250" align="left"></p>
        <p><br><img src="https://cdn-cndbm.nitrocdn.com/FELyUHTkveKxLHdFBEtBvbHkkGDKfZxM/assets/images/optimized/rev-6812291/site-content/uploads/2022/01/assorted-breakfast-meals-on-table-at-restaurant-sp.jpg"  height="300" align="left"></p>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">11/02 周四 可爱岛 一起上班 忙工作 看日落</h3>
        <p><br><img src="https://image.cnbcfm.com/api/v1/image/106845362-1614214121319-gettyimages-966275516-cr_soeg170104-00004-01.jpeg?v=1614214745"  height="250" align="right"></p>
        <p><br><img src="https://gowanderly.com/site-content/uploads/2023/02/Sunrise-over-the-coast-of-Kauai-Hawaii.jpg"  height="250" align="right"></p>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">11/03 周五 可爱岛 坐飞机 看日落</h3>
        <p><br><img src="https://media.tacdn.com/media/attractions-splice-spp-674x446/0b/d8/47/af.jpg"  height="250" align="left"></p>
                <p><br><img src="https://www.wanderingsunsets.com/wp-content/uploads/2017/09/FullSizeRender-6.jpg"  height="250" align="left"></p>

        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">11/04 周六 可爱岛 爬山 (跳伞) 抱一天</h3>
                <p><br><img src="https://scontent-sjc3-1.cdninstagram.com/v/t39.30808-6/336328467_544655537773767_7949993889363802673_n.jpg?stp=c0.60.1440.1800a_dst-jpg_e15&efg=eyJ2ZW5jb2RlX3RhZyI6ImltYWdlX3VybGdlbi4xNDQweDE5MjAuc2RyIn0&_nc_ht=scontent-sjc3-1.cdninstagram.com&_nc_cat=103&_nc_ohc=wpPa2yr4IA8AX-IILoh&edm=ABmJApAAAAAA&ccb=7-5&ig_cache_key=MzA2MDE3NTU4OTQ5NTA3NDEwMw%3D%3D.2-ccb7-5&oh=00_AfBeVTEPEtS5qTEjhers68r9Tzqhmhbb8hSPf_IPV_SPlw&oe=6525C0AE&_nc_sid=b41fef"  height="450" align="right"></p>
        <p><br><img src="https://tourscanner.com/blog/wp-content/uploads/2021/04/fun-things-to-do-in-Kauai-Hawaii.jpg"  height="250" align="right"></p>
        <p><br><img src="https://koloalandingresort.com/wp-content/uploads/2021/11/Romantic-couples-watching-sunset-at-the-beach.jpg"  height="250" align="right"></p>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title">11/05 周日 可爱岛 抱一天 回北京</h3>
        <p><br><img src="https://raw.githubusercontent.com/zzhangncsu/DataAnalysis/main/last.png"  height="250" align="left"></p>
        </div>
        </div>
        <div class="card">
        <div class="info">
        <h3 class="title"><a href="https://forms.office.com/Pages/ResponsePage.aspx?id=DQSIkWdsW0yxEjajBLZtrQAAAAAAAAAAAAIBAAI3sjNUMTFMOEFCTENENDM3RlNBM0dINzI5MVdSMy4u" target="_blank" >耍一哈？点击考虑一下<a></h3>
        </div>
        </div>
        </div>
        </div>
        </body>
            """,
            height=6000,
        )
    with tab5:
        url = st.text_input(
            "Enter URL"
        )
        if url:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_10_1) AppleWebKit/537.36 (KHTML, like Gecko) '
                              'Chrome/39.0.2171.95 Safari/537.36'}
            r = requests.get(url, headers=headers)
            soup = BeautifulSoup(r.text, "html.parser")
            sections = []
            for i in range(30):
                link = f"S-product-item j-expose__product-item recent_view_item_expose-{i} product-list__item"
                sections.append(soup.find_all("section", {"class": link})[0]["aria-label"])
            images = []
            for i in range(30):
                images.append(
                    "https://" + soup.find_all("div", {"class": "crop-image-container"})[i]["data-before-crop-src"].replace(
                        "//", ""))
            html = """
            <style>
            div.gallery {
              margin: 5px;
              border: 1px solid #ccc;
              float: left;
              width: 180px;
            }
            
            div.gallery:hover {
              border: 1px solid #777;
            }
            
            div.gallery img {
              width: 100%;
              height: auto;
            }
            
            div.desc {
              padding: 15px;
              text-align: center;
              color: white;
            }
            </style>
            
            """
            for img_url, text in zip(images, sections):
                html += f"""<div class ="gallery"> <img src="{img_url}" width="400"> <div class="desc">{text}</div></div>"""

            components.html(html,
            height=4000,
        )

